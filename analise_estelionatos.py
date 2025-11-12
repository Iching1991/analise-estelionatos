import pandas as pd
import matplotlib.pyplot as plt
import tkinter as tk
from tkinter import filedialog, simpledialog, messagebox
from docx import Document
import re
import string
import os

# ==============================
# FUNÇÕES AUXILIARES
# ==============================
def limpar_texto(texto):
    if isinstance(texto, str):
        texto = texto.encode('utf-8', errors='ignore').decode('utf-8', errors='ignore')
        texto = texto.lower().strip()
        texto = re.sub(r'\s+', ' ', texto)
        return texto
    return ""

def classificar_crime(texto):
    categorias = {
        "compras não autorizadas": [
            "compra", "cartão", "crédito", "débito", "não autoriz", "extrato", "fatura", "pagamento indevido"
        ],
        "golpe da maquininha": [
            "maquininha", "motoboy", "presente", "frete", "banco", "entregador", "cartão trocado"
        ],
        "golpe site falso / boleto falso": [
            "site", "boleto", "falso", "pix", "link", "página", "portal", "acesso", "conta"
        ],
        "whatsapp clonado": [
            "whatsapp", "mensagem", "clonado", "zap", "perfil falso", "pedido de dinheiro"
        ],
        "documentos clonados": [
            "documento", "rg", "cpf", "clonado", "identidade", "dados pessoais", "falsificado"
        ],
        "falsa central de atendimento": [
            "central", "atendimento", "ligação", "banco", "suporte", "falaram que era do banco", "funcionário"
        ],
        "golpe do namorado": [
            "namorado", "internet", "amor", "paixão", "relacionamento", "facebook", "instagram", "sedução"
        ],
        "golpe do carro": [
            "carro", "automóvel", "veículo", "revenda", "anúncio", "negócio", "venda"
        ],
        "golpe do bilhete premiado": [
            "bilhete", "premiado", "loteria", "sorteio", "prêmio", "idoso", "ajuda"
        ],
        "golpe do aluguel fantasma": [
            "aluguel", "imóvel", "casa", "apartamento", "kitnet", "corretor", "anúncio falso"
        ]
    }
    for categoria, palavras in categorias.items():
        if any(p in texto for p in palavras):
            return categoria
    return "Outros / Não identificado"

def coluna_para_indice(coluna_str):
    coluna_str = coluna_str.strip().upper()
    indice = 0
    for char in coluna_str:
        if char in string.ascii_uppercase:
            indice = indice * 26 + (ord(char) - ord('A') + 1)
    return indice - 1

# ==============================
# FUNÇÃO PRINCIPAL
# ==============================
def main():
    root = tk.Tk()
    root.withdraw()

    # Seleção do arquivo
    messagebox.showinfo("Seleção de Arquivo", "Selecione o arquivo Excel para análise.")
    caminho_arquivo = filedialog.askopenfilename(filetypes=[("Arquivos Excel", "*.xlsx")])
    if not caminho_arquivo:
        return

    df_teste = pd.read_excel(caminho_arquivo, header=None)
    total_linhas = len(df_teste)
    linha_inicio = simpledialog.askinteger("Linha Inicial", f"Digite a linha do cabeçalho (1 a {total_linhas}):", minvalue=1, maxvalue=total_linhas)
    if not linha_inicio:
        return
    df = pd.read_excel(caminho_arquivo, header=linha_inicio-1)

    # Colunas
    coluna_golpe_str = simpledialog.askstring("Coluna Golpes", "Digite a letra da coluna que contém os textos de golpes:")
    indice_golpe = coluna_para_indice(coluna_golpe_str)
    nome_coluna_golpe = df.columns[indice_golpe]

    coluna_bairro_str = simpledialog.askstring("Coluna Bairros", "Digite a letra da coluna com os bairros (ou deixe vazio):")
    if coluna_bairro_str:
        indice_bairro = coluna_para_indice(coluna_bairro_str)
        nome_coluna_bairro = df.columns[indice_bairro]
        df["Bairro Limpo"] = df[nome_coluna_bairro].astype(str).str.lower().str.strip()
    else:
        df["Bairro Limpo"] = "Não informado"

    # Coluna de idade (opcional)
    coluna_idade_str = simpledialog.askstring("Coluna Idade", "Digite a letra da coluna com a idade (ou deixe vazio):")
    idade_min = None
    if coluna_idade_str:
        indice_idade = coluna_para_indice(coluna_idade_str)
        nome_coluna_idade = df.columns[indice_idade]
        idade_min = simpledialog.askinteger("Filtro de Idade", "Mostrar apenas registros a partir de qual idade?", minvalue=0, maxvalue=120)

        # ✅ Correção segura da filtragem
        df[nome_coluna_idade] = pd.to_numeric(df[nome_coluna_idade], errors='coerce')
        df = df.dropna(subset=[nome_coluna_idade])
        df = df[df[nome_coluna_idade] >= idade_min]

    # Limpeza e classificação
    df["Texto Limpo"] = df[nome_coluna_golpe].apply(limpar_texto)
    df["Categoria"] = df["Texto Limpo"].apply(classificar_crime)

    # Contagem cruzada
    df_corr = df.groupby(["Categoria", "Bairro Limpo"]).size().unstack(fill_value=0)
    df_corr_percent = df_corr.div(df_corr.sum(axis=1), axis=0) * 100

    bairros = df_corr.columns.tolist()
    cores = plt.cm.tab20.colors
    cor_bairro_map = {bairro: cores[i % len(cores)] for i, bairro in enumerate(bairros)}

    # Pasta gráficos
    pasta_graficos = os.path.join(os.path.dirname(caminho_arquivo), "Graficos_Excel")
    os.makedirs(pasta_graficos, exist_ok=True)

    # ==============================
    # Gráfico de colunas e linhas
    # ==============================
    fig, ax = plt.subplots(figsize=(16, 8))
    df_corr_percent.plot(kind='bar', ax=ax, color=[cor_bairro_map[b] for b in bairros])

    for i, cat in enumerate(df_corr.index):
        for j, b in enumerate(bairros):
            v_abs = df_corr.loc[cat, b]
            if v_abs > 0:
                ax.text(i + (j - len(bairros)/2)*(0.8/len(bairros)), v_abs + 0.5, str(v_abs), ha='center', va='bottom', fontsize=8)

    totais_bairros = df.groupby('Bairro Limpo').size()
    legend_labels = [f"{bairro.title()} ({totais_bairros[bairro]})" for bairro in bairros]
    ax.legend(legend_labels, title="Bairros", bbox_to_anchor=(1.05, 1))
    ax.set_ylabel("Percentual (%)")
    ax.set_xlabel("Categoria de Golpe")
    ax.set_title("Distribuição Percentual de Golpes por Categoria e Bairro")
    ax.tick_params(axis='x', rotation=45)

    caminho_dashboard = os.path.join(pasta_graficos, "Dashboard_Resumo.png")
    plt.tight_layout()
    plt.savefig(caminho_dashboard)
    plt.close(fig)

    # ==============================
    # Salvar Excel
    # ==============================
    caminho_saida = caminho_arquivo.replace(".xlsx", "_analisado.xlsx")
    df.to_excel(caminho_saida, index=False)
    df_corr.to_excel(caminho_arquivo.replace(".xlsx", "_golpes_x_bairro.xlsx"))

    # ==============================
    # Word casos "Outros / Não identificado"
    # ==============================
    outros_df = df[df["Categoria"] == "Outros / Não identificado"][[nome_coluna_golpe, "Texto Limpo", "Bairro Limpo"]]
    if not outros_df.empty:
        doc = Document()
        doc.add_heading("Casos: Outros / Não Identificado", level=1)
        doc.add_paragraph(f"Total registros: {len(outros_df)}\n")
        for idx, linha in enumerate(outros_df.itertuples(), start=1):
            doc.add_paragraph(f"{idx}. Texto: {linha._2}\n   Bairro: {linha._3}")
        caminho_word = caminho_arquivo.replace(".xlsx", "_nao_identificados.docx")
        doc.save(caminho_word)
    else:
        caminho_word = None

    messagebox.showinfo("Concluído",
                        f"✅ Análise concluída!\n\nExcel: {caminho_saida}" +
                        (f"\nWord: {caminho_word}" if caminho_word else "") +
                        f"\nDashboard salvo em: {caminho_dashboard}")

# ==============================
# EXECUÇÃO
# ==============================
if __name__ == "__main__":
    main()
